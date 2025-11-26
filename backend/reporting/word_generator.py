from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
from datetime import datetime

class WordGenerator:
    def generate_analyst_report(self, results: dict, run_id: str) -> bytes:
        doc = Document()
        
        # Title Page
        title = doc.add_heading('Valuation Analyst Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        company_name = results.get('input_summary', {}).get('company_name', 'Target Company')
        doc.add_paragraph(f"Company: {company_name}")
        doc.add_paragraph(f"Date: {datetime.now().strftime('%B %d, %Y')}")
        doc.add_paragraph(f"Run ID: {run_id}")
        
        doc.add_page_break()
        
        # 1. Executive Summary
        doc.add_heading('1. Executive Summary', level=1)
        doc.add_paragraph(
            f"Based on the DCF analysis, the estimated Enterprise Value for {company_name} is "
            f"${results.get('enterprise_value', 0)/1000000:.1f} million. "
            f"The implied Equity Value is ${results.get('equity_value', 0)/1000000:.1f} million."
        )
        
        # Summary Table
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Light Shading Accent 1'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Metric'
        hdr_cells[1].text = 'Value'
        
        metrics = [
            ('Enterprise Value', f"${results.get('enterprise_value', 0):,.0f}"),
            ('Equity Value', f"${results.get('equity_value', 0):,.0f}"),
            ('WACC', f"{results.get('wacc', 0):.1%}")
        ]
        
        for metric, value in metrics:
            row_cells = table.add_row().cells
            row_cells[0].text = metric
            row_cells[1].text = value
            
        # 2. Methodology
        doc.add_heading('2. Methodology', level=1)
        doc.add_paragraph(
            "This valuation primarily relies on the Discounted Cash Flow (DCF) method. "
            "We project Free Cash Flow to Firm (FCFF) for a discrete period and calculate a Terminal Value "
            "using both the Gordon Growth Method and Exit Multiple method."
        )
        
        # 3. Detailed Projections
        doc.add_heading('3. Financial Projections', level=1)
        if 'dcf_details' in results:
            details = results['dcf_details']
            years = len(details.get('revenue', []))
            
            proj_table = doc.add_table(rows=1, cols=years + 1)
            proj_table.style = 'Table Grid'
            
            # Header
            hdr = proj_table.rows[0].cells
            hdr[0].text = 'Metric ($M)'
            for i in range(years):
                hdr[i+1].text = f"Year {i+1}"
                
            # Rows
            rows_data = [
                ('Revenue', details.get('revenue', [])),
                ('EBITDA', details.get('ebitda', [])),
                ('FCFF', details.get('fcff', []))
            ]
            
            for name, data in rows_data:
                row = proj_table.add_row().cells
                row[0].text = name
                for i, val in enumerate(data):
                    row[i+1].text = f"{val/1000000:.1f}"

        # 4. Sensitivity Analysis
        if 'sensitivity' in results and results['sensitivity']:
            doc.add_heading('4. Sensitivity Analysis', level=1)
            sens = results['sensitivity']
            doc.add_paragraph(f"Sensitivity of Enterprise Value to {sens['x_axis']['name']} and {sens['y_axis']['name']}.")
            
            # Matrix Table
            rows = len(sens['y_axis']['values'])
            cols = len(sens['x_axis']['values'])
            
            sens_table = doc.add_table(rows=rows + 1, cols=cols + 1)
            sens_table.style = 'Grid Table 4 Accent 1'
            
            # Header Row
            hdr = sens_table.rows[0].cells
            hdr[0].text = f"{sens['y_axis']['name']} \\ {sens['x_axis']['name']}"
            for i, val in enumerate(sens['x_axis']['values']):
                hdr[i+1].text = f"{val:.1%}"
                
            # Data Rows
            for i, row_val in enumerate(sens['y_axis']['values']):
                row = sens_table.rows[i+1].cells
                row[0].text = f"{row_val:.1%}"
                for j, val in enumerate(sens['matrix'][i]):
                    row[j+1].text = f"${val/1000000:.1f}M"

        # Save to buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.read()
